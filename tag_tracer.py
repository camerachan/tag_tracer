import re
import os
import openpyxl
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
import subprocess  # Excelファイルを開くために必要


class TagTracer:
    def __init__(self, root):
        self.root = root
        self.setup_gui()
    
    ##
    #@brief GUIをセットアップする
    ##
    def setup_gui(self):    
        self.root.title("Document Tag Tracer")
        self.root.geometry("700x150")
        ctk.set_appearance_mode("System") # Light/Dark/System
        ctk.set_default_color_theme("dark-blue")

        self.my_font = ("Meiryo UI", 12, "normal")

        # GUI部品のセットアップ
        # 初段
        self.label_file_a = ctk.CTkLabel(self.root, text = "トレース元文書", font=self.my_font)
        self.label_file_a.grid(row=0, column=0, padx=0, pady=5)

        self.entry_file_a = ctk.CTkEntry(self.root, width= 350)
        self.entry_file_a.grid(row=0, column=1, padx=0, pady=5)

        self.button_file_a = ctk.CTkButton(self.root, text="ファイル選択", command=self.select_file_a, font=self.my_font)
        self.button_file_a.grid(row=0, column=2, padx=0, pady=5)
        
        # 2段階目
        self.label_file_b = ctk.CTkLabel(self.root, text="トレース先文書", font=self.my_font)
        self.label_file_b.grid(row=1, column=0, padx=10, pady=5)

        self.entry_file_b = ctk.CTkEntry(self.root, width=350)
        self.entry_file_b.grid(row=1, column=1, padx=0, pady=5)

        self.button_file_b = ctk.CTkButton(self.root, text="ファイル選択", command=self.select_file_b, font=self.my_font)
        self.button_file_b.grid(row=1, column=2, padx=0, pady=5)

        # 3段階目
        #self.label_file_c = ctk.CTkLabel(self.root, text="多段トレース先文書", font=self.my_font)
        #self.label_file_c.grid(row=2, column=0, padx=10, pady=5)

        #self.entry_file_c = ctk.CTkEntry(self.root, width=350)
        #self.entry_file_c.grid(row=2, column=1, padx=0, pady=5)

        #self.button_file_c = ctk.CTkButton(self.root, text="ファイル選択", command=self.select_file_c, font=self.my_font)
        #self.button_file_c.grid(row=2, column=2, padx=0, pady=5)


        # 実行ボタン
        self.button_process = ctk.CTkButton(self.root, text="トレースファイル生成", command=self.process_files, font=self.my_font)
        self.button_process.grid(row=4, column=0, padx=2, pady=20)
        self.root.bind('<Return>', lambda event: self.process_files())

    
    ##
    # @brief トレース元ファイルを選択する関数
    ##
    def select_file_a(self):
        file_path = filedialog.askopenfilename(
            title="トレース元文書を選択",
            filetypes=[
                ("All Files", "*.*"),
                ("Text files", "*.txt")
            ]
        )
        if file_path:
            self.entry_file_a.delete(0, ctk.END)
            self.entry_file_a.insert(0, file_path)
    
    ##
    # @brief トレース先ファイルを選択する関数
    ##
    def select_file_b(self):  
        file_path = filedialog.askopenfilename(
            title="トレース先文書を選択",
            filetypes=[
                ("All Files", "*.*"),
                ("Text Files", "*.txt")
            ]
        )
        if file_path:
            self.entry_file_b.delete(0, ctk.END)
            self.entry_file_b.insert(0, file_path)

    ##
    # @brief トレース先ファイルを選択する関数
    ##
    def select_file_c(self):
        file_path = filedialog.askopenfilename(
            title="多段トレース文書を選択",
            filetypes=[
                ("All Files", "*.*"),
                ("Text Files", "*.txt")
            ]
        )
        if file_path:
            self.entry_file_c.delete(0, ctk.END)
            self.entry_file_c.insert(0, file_path)

    ##
    # @brief ドキュメントからタグを抽出する関数
    # @param document_text ドキュメントのテキスト内容
    # @return タグのリスト    
    ##
    def extract_tags(self, document_text):
        
        pattern = r'<[^<>]+>'
        tags = re.findall(pattern, document_text)
        print("抽出したタグ", tags)
        return tags
    ##
    # @brief タグのプレフィックス部分を抽出する関数
    # @param tags タグのリスト
    # @return プレフィックスのリスト
    ##    
    def extract_prefix(self, tags):  
        prefixes = set()
        for tag in tags:
            match = re.findall(r'<([A-Z0-9]+?)(?:\d+)>', tag)
            if match:
                prefixes.add(match[0])
        return list(prefixes)

    ##
    # @brief タグのプレフィックス部分を抽出する関数
    # @param tags タグのリスト
    # @return プレフィックスのリスト
    ##    
    def extract_target_prefix(self,tags):
        prefixes = set()
        for tag in tags:
            match = re.search(r'<([A-Z]+)(?=\d+-)', tag)
            if match:
                prefixes.add(match.group(1))
        
        return list(prefixes)

    ##
    # @brief ドキュメントからリンク情報を抽出する関数
    # @param doc_b トレース先ドキュメントの内容
    # @param tag_a トレース元タグのプレフィックス
    # @param tag_b トレース先タグのプレフィックス
    # @return リンク辞書とタグのリストのタプル
    ##
    def extract_links(self, doc_b, tag_a_list, tag_b_list):
        link_dict = {}
        b_tags_list = set()

        for tag_a in tag_a_list:
            for tag_b in tag_b_list:
                pattern = rf'\s*<\s*({re.escape(tag_b)}\d+)\s*[-–—]\s*({re.escape(tag_a)}\d+)\s*>\s*'
                links = re.findall(pattern, doc_b)
                print(f"抽出したトレースリンク ({tag_b} -> {tag_a}):", links)  # 抽出したリンクをコンソールに表示

                for b, a in links:
                    a_key = f"<{a}>"
                    b_tag = f"<{b}>"
                    b_tags_list.add(b_tag)

                    if a_key in link_dict:
                        link_dict[a_key].append(b_tag)
                    else:
                        link_dict[a_key] = [b_tag]

        return link_dict, sorted(b_tags_list)

    ##
    # @brief 多段トレースを実行する関数
    # @param doc_a トレース元ドキュメントの内容
    # @param doc_b 第1段階のトレース先ドキュメントの内容
    # @param doc_c 第2段階のトレース先ドキュメントの内容
    # @param tag_a_list トレース元タグのプレフィックスリスト
    # @param tag_b_list 第1段階のトレース先タグのプレフィックスリスト
    # @param tag_c_list 第2段階のトレース先タグのプレフィックスリスト
    ##
    def multi_level_link(self, doc_a, doc_b, doc_c, tag_a_list, tag_b_list, tag_c_list):
        links_a_b, b_tags_a_b = self.extract_links(doc_b, tag_a_list, tag_b_list)
        links_b_c, c_tags_b_c = self.extract_links(doc_c, tag_b_list, tag_c_list)
        multi_level_links = {}

        for a_tag, b_tags in links_a_b.items():
            for b_tag in b_tags:
                if b_tag in links_b_c:
                    for c_tag in links_b_c[b_tag]:
                        if a_tag not in multi_level_links:
                            multi_level_links[a_tag] = {}
                        if b_tag not in multi_level_links:
                            multi_level_links[a_tag][b_tag] = []
                        multi_level_links[a_tag][b_tag].append(c_tag)
        return multi_level_links        


    ##
    # @brief ファイルの内容を読み取る関数
    # @param file_path ファイルのパス
    # @return ファイルの内容
    ##
    def read_document(self, file_path):
        
        ext = os.path.splitext(file_path)[1].lower()

        if ext in [".txt", ".c", ".cpp", ".h", ".hpp"]:
            # テキストファイル、C/C++ファイルの場合
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif ext == ".docx":
            # Wordファイルの場合
            return self.read_docx_full(file_path)
        elif ext == ".xlsx":
            # Excelファイルの場合
            return self.read_xlsx(file_path)
        else:
            raise ValueError("サポートされていないファイル形式です")
    ##
    # @brief DOCXファイルを読み込む関数
    # @param file_path ファイルのパス
    # @return ファイルの内容
    ##
    def read_docx(self, file_path):
        doc = Document(file_path)
        # 全ての段落を結合して単一のテキストとして取得
        doc_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        return doc_text
    ##
    # @brief DOCXファイル全体からテキストを抽出する関数
    # @param file_path ファイルのパス
    # @return ファイルの内容
    ##
    def read_docx_full(self, file_path):
        
        doc = Document(file_path)
        # すべてのテキストを収集
        full_text = []

        # 各段落のテキストを収集
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 各テーブルのテキストを収集
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)

        # 全体を1つのテキストに結合
        doc_text = "\n".join(full_text)
        print("全体のテキスト:\n", doc_text)
        return doc_text
    
    ##
    # @brief XLSXファイルを読み込む関数
    # @param file_path ファイルのパス
    # @return ファイルの内容
    ##
    def read_xlsx(self, file_path):

        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        excel_text = ""
    
        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join([str(cell) for cell in row if cell is not None])
            excel_text += row_text + "\n"

        return excel_text

    ##
    # @brief タグトレース結果をExcelファイルに出力する関数
    # @param all_tags_a トレース元のタグリスト
    # @param links リンクの辞書
    # @param b_tags トレース先のタグリスト
    # @param output_file 出力ファイル名
    # @param input_file1 トレース元ファイル名
    # @param input_file2 トレース先ファイル名
    ##
    def create_excel_file(self, all_tags_a, links, b_tags, output_file, input_file1, input_file2):
        
        input_file1_name = os.path.basename(input_file1)
        input_file2_name = os.path.basename(input_file2)

        wb = openpyxl.Workbook()

        # 1つ目のシート: トレース結果
        ws1 = wb.active
        ws1.title = "Traceability"

        ws1["A1"] = 'トレース元'
        ws1["B1"] = 'トレース先'
        ws1["A2"] = input_file1_name
        ws1["B2"] = input_file2_name

        # 文書Aの全タグを出力
        for idx, a_tag in enumerate(all_tags_a, start=3):
            ws1[f'A{idx}'] = a_tag
            # 文書Aのタグで対応する文書Bのタグをリストで取得
            b_tags_list = links.get(a_tag, [])
            if b_tags_list:
                # 複数のBタグを結合して1つのセルに表示
                ws1[f'B{idx}'] = ', '.join(b_tags_list)
            else:
                ws1[f'B{idx}'] = ""  # 見つからない場合は空白

        # 2つ目のシート: 星取表
        ws2 = wb.create_sheet(title="Star Chart")

        # 文書Bのタグを横軸に設定
        ws2["A1"] = 'トレース元'
        for col_idx, b_tag in enumerate(b_tags, start=2):
            ws2.cell(row=1, column=col_idx, value=b_tag)

        # 文書Aのタグを縦軸に設定し、対応するBタグに●を記入
        for row_idx, a_tag in enumerate(all_tags_a, start=2):
            ws2.cell(row=row_idx, column=1, value=a_tag)  # 文書AのタグをA列に設定
            b_tags_list = links.get(a_tag, [])
            for b_tag in b_tags_list:
                if b_tag in b_tags:
                    col_idx = b_tags.index(b_tag) + 2  # 横軸に対応するBタグのインデックスを取得
                    ws2.cell(row=row_idx, column=col_idx, value="●")  # リンクされているところに●を記入

        # Excelファイルを保存
        wb.save(output_file)
        print(f"Excel File '{output_file}' has created.")
        #messagebox.showinfo("Success", f"Excel File '{output_file}' has created.")

    ##
    # @brief 多段トレース結果をExcelファイルに出力する関数
    # @param multi_level_links 多段リンク情報の辞書
    # @param output_file 出力するExcelファイルのパス
    # @param input_file1 トレース元ファイル名
    # @param input_file2 第1段階のトレース先ファイル名
    # @param input_file3 第2段階のトレース先ファイル名
    ##
    def create_excel_file_multi(self, multi_level_links, output_file, input_file1, input_file2, input_file3):
        input_file1_name = os.path.basename(input_file1)
        input_file2_name = os.path.basename(input_file2)
        input_file3_name = os.path.basename(input_file3)

        wb = openpyxl.Workbook()

        ws1 = wb.active
        ws1.title = "Traceability"

        ws1["A1"] = 'トレース元'
        ws1["B1"] = '第1段階トレース先'
        ws1["C1"] = '第2段階トレース先'
        ws1["A2"] = input_file1_name
        ws1["B2"] = input_file2_name
        ws1["C2"] = input_file3_name

        row = 3
        
        for a_tag, b_data in multi_level_links.items():
            for b_tag, c_tags in b_data.items():
                for c_tag in c_tags:
                    ws1.cell(row=row, column=1, value=a_tag)
                    ws1.cell(row=row, column=2, value=b_tag)
                    ws1.cell(row=row, column=3, value=c_tag)
                    row += 1
        ws2 = wb.create_sheet(title="Star Chart")

        ws2["A1"] = 'トレース元'
        b_tag_set = {b_tag for b_data in multi_level_links.values() for b_tag in b_data}
        c_tag_set = {c_tag for b_data in multi_level_links.values() for c_tags in b_data.values() for c_tag in c_tags}
        b_tag_list = sorted(b_tag_set)
        c_tag_list = sorted(c_tag_set)

        for col_idx, b_tag in enumerate(b_tag_list, start=2):
            ws2.cell(row=1, column=col_idx, value=b_tag)
        for row_idx, c_tag in enumerate(c_tag_list, start=2):
            ws2.cell(row=row_idx, column=1, value=c_tag)

        row_start = 2
        for a_tag, b_data in multi_level_links.items():
            ws2.cell(row=row_start, column=1, value=a_tag)
            for b_tag in b_data:
                b_col = b_tag_list.index(b_tag) + 2
                ws2.cell(row=row_start, column=b_col, value="●")
                for c_tag in b_data[b_tag]:
                    c_row = c_tag_list.index(c_tag) + 2
                    ws2.cell(row=c_row, column=b_col, value="●")
            row_start += 1

        wb.save(output_file)


    ##
    # @brief トレースファイルの生成プロセスを実行する関数
    ##
    def process_files(self):
        
        file_a = self.entry_file_a.get()
        file_b = self.entry_file_b.get()
        file_c = self.entry_file_c.get()
    
        if not file_a or not file_b:
            messagebox.showwarning("Warning", "すべてのフィールドに入力してください")
            return
        # トレース先、トレース元の文書を読み込む
        doc_a = self.read_document(file_a)
        doc_b = self.read_document(file_b)
        # 文書Aから全てのタグを抽出
        all_tags_a = self.extract_tags(doc_a)
        tag_a = self.extract_prefix(all_tags_a)
        print(tag_a)
        all_tags_b = self.extract_tags(doc_b)
        tag_b = self.extract_target_prefix(all_tags_b)
        print(tag_b)
        if not file_c:
            # 文書Bからリンクを抽出し、Bの全タグリストを取得
            links, b_tags = self.extract_links(doc_b, tag_a, tag_b)

            output_file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if output_file:
                # Excelファイルを生成
                self.create_excel_file(all_tags_a, links, b_tags, output_file, file_a, file_b)
                
                # Excelファイルを自動で開く
                try:    
                    os.startfile(output_file)
                except Exception as e:
                    messagebox.showerror("Error", f"ファイルを開く際にエラーが発生しました: {str(e)}")
        else:
            doc_c = self.read_document(file_c)
            all_tags_c = self.extract_tags(doc_c)
            tag_c = self.extract_target_prefix(all_tags_c)
            print(tag_c)
            multi_level_links = self.multi_level_link(doc_a, doc_b, doc_c, tag_a, tag_b, tag_c)
            output_file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if output_file:
                self.create_excel_file_multi(multi_level_links, output_file, file_a, file_b, file_c)
                try:
                    os.startfile(output_file)
                except Exception as e:
                    messagebox.showerror("Error", f"ファイルを開く際にエラーが発生しました: {str(e)}")
            

if __name__ == "__main__":
    root = ctk.CTk()
    app = TagTracer(root)
    root.mainloop()